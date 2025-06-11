import os
import io
from PyPDF2 import PdfReader # PyPDF2 is the successor to PyPDF; ensure it's in requirements.txt
import docx # python-docx library

from job_application_agent.core_modules.error_handler import CVParserError, get_logger

logger = get_logger(__name__)

def _extract_text_from_pdf(file_stream: io.BytesIO) -> str:
    """Extracts text from a PDF file stream."""
    try:
        pdf_reader = PdfReader(file_stream)
        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                text_parts.append(page.extract_text() or "") # Ensure None is handled
            except Exception as e:
                logger.warning(f"Could not extract text from page {page_num + 1} of PDF: {e}")
                text_parts.append("") # Add empty string for pages that fail extraction
        full_text = "\n".join(text_parts)
        if not full_text.strip():
            logger.warning("PDF text extraction resulted in empty or whitespace-only content.")
        return full_text
    except Exception as e:
        logger.error(f"Failed to read or parse PDF stream: {e}", exc_info=True)
        raise CVParserError(f"Error processing PDF file: {e}")


def _extract_text_from_docx(file_stream: io.BytesIO) -> str:
    """Extracts text from a DOCX file stream."""
    try:
        document = docx.Document(file_stream)
        text_parts = [paragraph.text for paragraph in document.paragraphs]
        full_text = "\n".join(text_parts)
        if not full_text.strip():
            logger.warning("DOCX text extraction resulted in empty or whitespace-only content.")
        return full_text
    except Exception as e:
        logger.error(f"Failed to read or parse DOCX stream: {e}", exc_info=True)
        # python-docx can raise various errors, e.g., PackageNotFoundError for corrupted files
        raise CVParserError(f"Error processing DOCX file: {e}")


def parse_cv(file_path_or_stream, file_name: str) -> str:
    """
    Parses a CV file (PDF or DOCX) and extracts its text content.

    Args:
        file_path_or_stream (str or io.BytesIO):
            Either a string representing the full path to the CV file,
            or a file-like object (e.g., io.BytesIO) containing the CV data.
        file_name (str): The original name of the file, used to determine file type.

    Returns:
        str: The extracted raw text content from the CV.

    Raises:
        CVParserError: If the file type is unsupported, if parsing fails,
                       or if the input type is invalid.
        FileNotFoundError: If file_path_or_stream is a path and the file is not found.
    """
    logger.info(f"Attempting to parse CV: {file_name}")

    _, file_extension = os.path.splitext(file_name.lower())

    text_content = ""

    try:
        if isinstance(file_path_or_stream, str): # It's a file path
            logger.debug(f"Parsing CV from file path: {file_path_or_stream}")
            if not os.path.exists(file_path_or_stream):
                logger.error(f"File not found at path: {file_path_or_stream}")
                raise FileNotFoundError(f"CV file not found: {file_path_or_stream}")

            with open(file_path_or_stream, 'rb') as f_stream:
                if file_extension == '.pdf':
                    text_content = _extract_text_from_pdf(f_stream)
                elif file_extension == '.docx':
                    text_content = _extract_text_from_docx(f_stream)
                else:
                    logger.warning(f"Unsupported file type: {file_extension} for file {file_name}")
                    raise CVParserError(f"Unsupported file type: {file_extension}. Please upload a PDF or DOCX file.")

        elif hasattr(file_path_or_stream, 'read'): # It's a file-like object (stream)
            logger.debug(f"Parsing CV from stream: {file_name}")
            # Ensure it's a BytesIO stream if it's not already
            if not isinstance(file_path_or_stream, io.BytesIO):
                # If it's another type of stream (e.g., SpooledTemporaryFile from Telegram),
                # read its content into BytesIO.
                # This might not be strictly necessary if PyPDF2/python-docx can handle other stream types,
                # but BytesIO is generally safe.
                current_pos = file_path_or_stream.tell() # Remember current position
                file_path_or_stream.seek(0)
                file_bytes = file_path_or_stream.read()
                file_path_or_stream.seek(current_pos) # Reset position for other potential readers

                stream_to_parse = io.BytesIO(file_bytes)
            else:
                stream_to_parse = file_path_or_stream

            # Ensure stream is at the beginning before parsing
            stream_to_parse.seek(0)

            if file_extension == '.pdf':
                text_content = _extract_text_from_pdf(stream_to_parse)
            elif file_extension == '.docx':
                text_content = _extract_text_from_docx(stream_to_parse)
            else:
                logger.warning(f"Unsupported file type: {file_extension} for file {file_name}")
                raise CVParserError(f"Unsupported file type: {file_extension}. Please upload a PDF or DOCX file.")

            stream_to_parse.seek(0) # Reset stream position again after reading, good practice

        else:
            logger.error(f"Invalid input type for file_path_or_stream: {type(file_path_or_stream)}")
            raise CVParserError("Invalid input: Must be a file path string or a file-like stream object.")

        if not text_content.strip():
            logger.warning(f"Parsing of '{file_name}' resulted in empty text content. The file might be image-based or corrupted.")
            # Not raising an error here, as an empty text is a valid parsing outcome for some files (e.g. scanned PDF)
            # The LLM module will have to handle empty text if it receives it.

        logger.info(f"Successfully parsed CV: {file_name}. Extracted text length: {len(text_content)}")
        return text_content.strip() # Return stripped text

    except FileNotFoundError as fnf_err: # Specifically re-raise FileNotFoundError
        raise fnf_err
    except CVParserError: # Re-raise CVParserError directly
        raise
    except Exception as e: # Catch any other unexpected errors during parsing
        logger.error(f"An unexpected error occurred while parsing {file_name}: {e}", exc_info=True)
        raise CVParserError(f"An unexpected error occurred while processing {file_name}: {e}")


# --- Example Usage (for testing) ---
if __name__ == '__main__':
    # This requires actual PDF and DOCX files to be present in a 'test_cvs' directory
    # relative to where this script is run, or you need to adjust paths.
    # For demonstration, we'll create dummy files if they don't exist.

    # Setup basic logging for testing this module standalone
    import logging
    logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')

    # Create a dummy test_cvs directory if it doesn't exist
    test_cv_dir = "test_cvs_for_parser"
    if not os.path.exists(test_cv_dir):
        os.makedirs(test_cv_dir)

    # Dummy PDF file creation
    dummy_pdf_path = os.path.join(test_cv_dir, "dummy_cv.pdf")
    if not os.path.exists(dummy_pdf_path):
        try:
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(dummy_pdf_path)
            c.drawString(100, 750, "This is a dummy PDF CV.")
            c.drawString(100, 730, "John Doe - john.doe@example.com")
            c.drawString(100, 710, "Skills: Python, ReportLab.")
            c.save()
            logger.info(f"Created dummy PDF: {dummy_pdf_path}")
        except ImportError:
            logger.warning("ReportLab not installed. Cannot create dummy PDF for testing. Please create one manually.")
        except Exception as e:
            logger.warning(f"Could not create dummy PDF: {e}")


    # Dummy DOCX file creation
    dummy_docx_path = os.path.join(test_cv_dir, "dummy_cv.docx")
    if not os.path.exists(dummy_docx_path):
        try:
            doc = docx.Document()
            doc.add_paragraph("This is a dummy DOCX CV.")
            doc.add_paragraph("Jane Doe - jane.doe@example.com")
            doc.add_paragraph("Skills: Python, python-docx.")
            doc.save(dummy_docx_path)
            logger.info(f"Created dummy DOCX: {dummy_docx_path}")
        except Exception as e:
            logger.warning(f"Could not create dummy DOCX: {e}")

    # Test with file paths
    if os.path.exists(dummy_pdf_path):
        print(f"\n--- Testing PDF from path: {dummy_pdf_path} ---")
        try:
            pdf_text_path = parse_cv(dummy_pdf_path, "dummy_cv.pdf")
            print(f"Extracted PDF text (path):\n{pdf_text_path}\n")
        except Exception as e:
            print(f"Error parsing PDF from path: {e}")

    if os.path.exists(dummy_docx_path):
        print(f"\n--- Testing DOCX from path: {dummy_docx_path} ---")
        try:
            docx_text_path = parse_cv(dummy_docx_path, "dummy_cv.docx")
            print(f"Extracted DOCX text (path):\n{docx_text_path}\n")
        except Exception as e:
            print(f"Error parsing DOCX from path: {e}")

    # Test with streams
    if os.path.exists(dummy_pdf_path):
        print(f"\n--- Testing PDF from stream ---")
        try:
            with open(dummy_pdf_path, 'rb') as f:
                pdf_stream = io.BytesIO(f.read())
            pdf_text_stream = parse_cv(pdf_stream, "dummy_cv_stream.pdf")
            print(f"Extracted PDF text (stream):\n{pdf_text_stream}\n")
        except Exception as e:
            print(f"Error parsing PDF from stream: {e}")

    if os.path.exists(dummy_docx_path):
        print(f"\n--- Testing DOCX from stream ---")
        try:
            with open(dummy_docx_path, 'rb') as f:
                docx_stream = io.BytesIO(f.read())
            docx_text_stream = parse_cv(docx_stream, "dummy_cv_stream.docx")
            print(f"Extracted DOCX text (stream):\n{docx_text_stream}\n")
        except Exception as e:
            print(f"Error parsing DOCX from stream: {e}")

    # Test unsupported file type
    print("\n--- Testing unsupported file type ---")
    try:
        unsupported_text = parse_cv(io.BytesIO(b"dummy content"), "dummy_cv.txt")
        print(f"Unsupported file text (should not happen):\n{unsupported_text}\n")
    except CVParserError as e:
        print(f"Caught expected error for unsupported file: {e}\n")
    except Exception as e:
        print(f"Caught unexpected error for unsupported file: {e}\n")

    # Test non-existent file
    print("\n--- Testing non-existent file ---")
    try:
        parse_cv("non_existent_cv.pdf", "non_existent_cv.pdf")
    except FileNotFoundError as e:
        print(f"Caught expected FileNotFoundError: {e}\n")
    except Exception as e:
        print(f"Caught unexpected error: {e}\n")

    # Test with a potentially problematic PDF (e.g. image-based, if you have one)
    # Create an empty PDF (often problematic for some parsers if not handled)
    empty_pdf_path = os.path.join(test_cv_dir, "empty.pdf")
    if not os.path.exists(empty_pdf_path):
        try:
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(empty_pdf_path)
            c.save()
            logger.info(f"Created empty PDF: {empty_pdf_path}")
        except ImportError:
            logger.warning("ReportLab not installed. Cannot create empty PDF for testing.")
        except Exception as e:
            logger.warning(f"Could not create empty PDF: {e}")

    if os.path.exists(empty_pdf_path):
        print(f"\n--- Testing Empty PDF from path: {empty_pdf_path} ---")
        try:
            empty_pdf_text = parse_cv(empty_pdf_path, "empty.pdf")
            print(f"Extracted Empty PDF text (path should be empty or warning issued):\n'{empty_pdf_text}'\n")
        except Exception as e:
            print(f"Error parsing Empty PDF from path: {e}")

    print("--- CV Parser standalone test complete ---")
